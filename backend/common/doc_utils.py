from django.utils.text import capfirst

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Inches


class DocumentService:
    @classmethod
    def safe_get(cls, value):
        return value if value not in [None, '', []] else "N/A"

    @classmethod
    def add_section(cls, document, title, data_dict):
        document.add_heading(title, level=1)
        for key, value in data_dict.items():
            document.add_heading(capfirst(key.replace("_", " ")), level=2)
            document.add_paragraph(cls.safe_get(value))

    @classmethod
    def add_table(cls, document, title, data_dict):
        document.add_heading(title, level=1)
        table = document.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        for key, value in data_dict.items():
            row_cells = table.add_row().cells
            row_cells[0].text = capfirst(key.replace("_", " "))
            row_cells[1].text = cls.safe_get(value)

            # Apply padding to both cells in the row
            for cell in row_cells:
                cls.set_cell_margins(cell, top=100, start=100, bottom=100, end=100)

    @staticmethod
    def set_cell_margins(cell, **kwargs):
        """
        Set cell margins (padding) in twips: 1 twip = 1/20 pt = 1/1440 inch.
        Example: 100 = ~5pt of padding.
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m in ['top', 'start', 'bottom', 'end']:
            if m in kwargs:
                node = OxmlElement(f'w:{m}')
                node.set(qn('w:w'), str(kwargs[m]))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
        tcPr.append(tcMar)

    @classmethod
    def set_cell_background(cls, cell, fill_color_hex):
        """
        Set the background color of a cell.
        Args:
            cell: The cell object.
            fill_color_hex: A hex color string like '0044CC' (without the #).
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')  # clear means solid fill
        shd.set(qn('w:color'), 'auto')  # font color
        shd.set(qn('w:fill'), fill_color_hex)  # background fill
        tcPr.append(shd)
    
    @classmethod
    def set_cell_font_color(cls, cell, text, hex_color="000000", bold=False):
        """
        Set the font color and bold style for a cell's text.
        
        Args:
            cell: The cell object from a table row.
            text: The string content to add.
            hex_color: Font color in hex (e.g., '004B8D').
            bold: Boolean indicating if text should be bold.
        """
        # Convert hex to RGB
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)

        # Clear existing text
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = bold
        run.font.color.rgb = RGBColor(r, g, b)

    @classmethod
    def css_color_to_rgb_color(cls, css_color):
        if css_color.startswith("#"):
            css_color = css_color.lstrip("#")
            return RGBColor(int(css_color[0:2], 16), int(css_color[2:4], 16), int(css_color[4:6], 16))
        elif css_color.startswith("rgb"):
            nums = [int(n) for n in css_color[4:-1].split(",")]
            return RGBColor(*nums)
        else:
            raise ValueError("Unsupported color format. Use '#rrggbb' or 'rgb(r, g, b)'")
    
    @classmethod
    def add_table_from_list(cls, document, title, data_list, headers=None, headers_color_config=None):
        """
        Add a table to the document using a list of lists or list of dicts.

        Args:
            document: The docx.Document object.
            title (str): The title for the section.
            data_list (list): List of rows (each row is a list or dict).
            headers (list or None): If provided, creates a header row using these values.
            headers_color_config (dict): Optional color config for header cells.
        """
        if title:
            document.add_heading(title, level=1)

        if not data_list:
            document.add_paragraph("No data available.")
            return

        if isinstance(data_list[0], dict):
            if headers is None:
                headers = list(data_list[0].keys())
            rows = [[str(row.get(h, '')) for h in headers] for row in data_list]
        else:
            rows = data_list

        table = document.add_table(rows=0, cols=len(headers) if headers else len(rows[0]))
        table.style = 'Table Grid'

        # Add header row
        if headers:
            hdr_cells = table.add_row().cells
            for i, header in enumerate(headers):
                header_text = capfirst(str(header))
                hdr_cells[i].text = header_text

                # Set font and background colors
                if headers_color_config:
                    cls.set_cell_background(hdr_cells[i], headers_color_config.get("background"))
                    cls.set_cell_font_color(hdr_cells[i], header_text, headers_color_config.get("font_color"))

                cls.set_cell_margins(hdr_cells[i], top=100, start=100, bottom=100, end=100)

        # Add data rows
        for row in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                cls.set_cell_margins(row_cells[i], top=100, start=100, bottom=100, end=100)

    @classmethod
    def add_bold_label(cls, document, label, value):
        p = document.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))


